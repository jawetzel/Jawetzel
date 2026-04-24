"""
Builds the redacted security-audit .docx — an intermediate artifact used
only to produce the final PDF deliverable.

Reads the four sanitized markdown files in ./_content/ and emits
./_build/Security_Audit_Report_Redacted.docx. Neither the .docx nor the
_build/ folder is served — they are not deliverables.

To produce the downloadable PDF: open the .docx in Word (or any
equivalent) and Save As PDF into
public/security-audit/Security_Audit_Report_Redacted.pdf — that is the
path the page's download link expects.

Usage (from repo root):
    python src/app/security-audit/build_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

HERE = Path(__file__).resolve().parent
CONTENT_DIR = HERE / "_content"
OUTPUT_DIR = HERE / "_build"
OUTPUT_FILE = OUTPUT_DIR / "Security_Audit_Report_Redacted.docx"


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg: str = "1F2937", fg: str = "FFFFFF") -> None:
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string(fg)
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    style_header_row(table.rows[0])

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            val = row_data[c_idx] if c_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, val, size=Pt(9))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F4F6")


# ---------------------------------------------------------------------------
# Minimal inline-markdown renderer (bold, italic, inline code, links)
# ---------------------------------------------------------------------------

_INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def _add_inline_runs(paragraph, text: str, size: Pt | None = None) -> None:
    """Render inline markdown (bold/italic/code/links) into a paragraph."""
    for part in _INLINE_TOKEN.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        if size is not None:
            run.font.size = size
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            if size is None:
                run.font.size = Pt(9.5)
        elif part.startswith("[") and "](" in part:
            label, _ = part[1:-1].split("](", 1)
            run.text = label
        else:
            run.text = part


# ---------------------------------------------------------------------------
# Markdown block parser — handles what our redacted files actually use
# ---------------------------------------------------------------------------

def strip_docx_skip_blocks(md: str) -> str:
    """Remove any `<!-- docx:skip-block -->...<!-- /docx:skip-block -->`
    chunks so they never reach the PDF. Use for trimming a paragraph,
    table, or bullet group inside a section you otherwise want to keep."""
    return re.sub(
        r"<!--\s*docx:skip-block\s*-->.*?<!--\s*/docx:skip-block\s*-->\s*",
        "",
        md,
        flags=re.DOTALL,
    )


def strip_docx_skips(md: str) -> str:
    """Remove sections marked `<!-- docx:skip -->` so they never reach the PDF.

    The marker must appear on its own line immediately after a heading
    (blank lines allowed between). The heading and every line up to the
    next heading of same-or-higher level are dropped.
    """
    lines = md.splitlines()
    result: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        h = re.match(r"^(#{1,6})\s+", line)
        if h:
            level = len(h.group(1))
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].strip() == "<!-- docx:skip -->":
                i += 1
                while i < len(lines):
                    m2 = re.match(r"^(#{1,6})\s+", lines[i])
                    if m2 and len(m2.group(1)) <= level:
                        break
                    i += 1
                continue
        result.append(line)
        i += 1
    return "\n".join(result)


def parse_blocks(md: str) -> list[tuple]:
    """Return a list of (kind, payload) tuples describing the document."""
    lines = md.splitlines()
    blocks: list[tuple] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---"):
            blocks.append(("hr", None))
            i += 1
            continue

        if re.fullmatch(r"<!--\s*pagebreak\s*-->", stripped):
            blocks.append(("pagebreak", None))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append(("heading", (level, m.group(2).strip())))
            i += 1
            continue

        # Markdown table: header line + separator + rows
        if "|" in stripped and i + 1 < len(lines) and re.match(
            r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1]
        ):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and "|" in lines[j].strip():
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            blocks.append(("table", (header, rows)))
            i = j
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            items: list[str] = []
            while i < len(lines) and (
                lines[i].strip().startswith("- ")
                or lines[i].strip().startswith("* ")
            ):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("bullets", items))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            num_items: list[str] = []
            while i < len(lines):
                m = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if not m:
                    break
                num_items.append(m.group(1))
                i += 1
            blocks.append(("numbered", num_items))
            continue

        # Paragraph — collect consecutive non-empty, non-table, non-list lines
        para: list[str] = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if nxt.startswith("#") or nxt.startswith("---"):
                break
            if nxt.startswith("- ") or nxt.startswith("* "):
                break
            if "|" in nxt and j + 1 < len(lines) and re.match(
                r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[j + 1]
            ):
                break
            para.append(nxt)
            j += 1
        blocks.append(("paragraph", " ".join(para)))
        i = j

    return blocks


# ---------------------------------------------------------------------------
# Document emitter
# ---------------------------------------------------------------------------

def emit_blocks(doc, blocks: list[tuple], heading_base: int = 1) -> None:
    """Render parsed blocks into a docx Document, shifting heading levels
    down by (heading_base - 1) so a section can be embedded under a higher-
    level heading."""
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            shifted = min(9, level + heading_base - 1)
            doc.add_heading(text, level=shifted)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            _add_inline_runs(p, payload)
        elif kind == "bullets":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item)
        elif kind == "numbered":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item)
        elif kind == "table":
            header, rows = payload
            add_table(doc, header, rows)
            doc.add_paragraph("")
        elif kind == "hr":
            # Treat horizontal rules as light section separators.
            doc.add_paragraph("")
        elif kind == "pagebreak":
            doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def build_doc() -> Document:
    doc = Document()

    # Default body font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # Heading styles
    for level, size, color in [(1, 18, "1F2937"), (2, 14, "1F2937"), (3, 12, "374151")]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(size)
        hs.font.color.rgb = RGBColor.from_string(color)
        hs.font.bold = True

    # Narrow margins for table-heavy pages
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _add_title_page(doc)
    _add_toc(doc)
    _add_section(doc, "summary.md", heading="1. Executive Summary")
    doc.add_page_break()
    _add_section(doc, "details-dashboards.md", heading="2. Internal Pages Exposed — Detail")
    doc.add_page_break()
    _add_section(doc, "details-data-leaks.md", heading="3. Data Leaks — Detail")
    doc.add_page_break()
    _add_section(doc, "scope.md", heading="4. What We Looked At")
    _add_footer(doc)

    return doc


def _add_title_page(doc) -> None:
    for _ in range(6):
        doc.add_paragraph("")

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run("SECURITY AUDIT REPORT")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string("1F2937")
    r.bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("[company] — redacted portfolio version")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_paragraph("")

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Non-Intrusive Public Information Review")
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor.from_string("6B7280")

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("Early 2026")
    r4.font.size = Pt(12)
    r4.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_page_break()


def _add_toc(doc) -> None:
    doc.add_heading("Table of Contents", level=1)
    for item in [
        "1. Executive Summary",
        "2. Internal Pages Exposed — Detail",
        "3. Data Leaks — Detail",
        "4. What We Looked At",
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def _add_section(doc, filename: str, heading: str) -> None:
    md = (CONTENT_DIR / filename).read_text(encoding="utf-8")
    md = strip_docx_skip_blocks(md)
    md = strip_docx_skips(md)
    # Strip the first top-level heading; we replace it with our numbered one.
    md = re.sub(r"^\s*#\s+.*\n", "", md, count=1)
    doc.add_heading(heading, level=1)
    blocks = parse_blocks(md)
    emit_blocks(doc, blocks, heading_base=2)


def _add_footer(doc) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r = p.add_run("Security Audit Report — redacted portfolio version  |  Page ")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("9CA3AF")

        fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        p._p.append(fld_begin)
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        p._p.append(instr)
        fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        p._p.append(fld_end)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
